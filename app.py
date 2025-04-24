import streamlit as st
from openai import OpenAI
from docx import Document
import fitz
import csv
from datetime import datetime
import pandas as pd
import io
import re

# 👉 First Streamlit command
st.set_page_config(page_title="Career Coach AI", page_icon="🧠")

# Initialize session state
if "started" not in st.session_state:
    st.session_state.started = False

# 🎨 Landing view
if not st.session_state.started:
    hero_url = (
        "https://raw.githubusercontent.com/NicholasRico/"
        "career-coach-ai/main/.streamlit/assets/career-coach-hero-ng.jpg"
    )
    st.title("Career Coach AI")
    st.image(hero_url, use_container_width=True)
    st.markdown("Tailor your resume, cover letter, and recruiter message for **any job** in seconds.")
    st.markdown("Built by [Nicholas Gauthier](mailto:NickRGauthier@gmail.com)")
    if st.button("🚀 Get Started"):
        st.session_state.started = True
    st.stop()

# 🎨 Content background CSS (only after landing)
content_bg = (
    "https://raw.githubusercontent.com/NicholasRico/"
    "career-coach-ai/main/.streamlit/assets/career-coach-content-ng.jpg"
)
st.markdown(
    f"""
    <style>
    .stApp {{
        background-image: url('{content_bg}');
        background-size: cover;
        background-repeat: no-repeat;
        background-attachment: fixed;
        background-position: center top;
    }}
    </style>
    """,
    unsafe_allow_html=True,
)

# 👤 Track who is using it (normalized email/name)
user_id = st.text_input("👤 Enter your name or email to track your applications").strip().lower()
if not user_id:
    st.warning("Please enter your name or email to continue.")
    st.stop()

# 🔌 OpenAI client
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

# --- Main Inputs + Quick Steps ---
st.markdown("---")
st.markdown("## Job Application Tailoring")
col_main, col_help = st.columns([3,1])

with col_main:
    resume_file = st.file_uploader("📌 Upload your resume (.pdf or .docx)", type=["pdf","docx"])
    job_desc    = st.text_area("📜 Paste job description here", height=250)
    model       = st.selectbox("🧠 Choose GPT model", ["gpt-3.5-turbo","gpt-4"])

    c1,c2,c3 = st.columns(3)
    with c1:
        log_app = st.checkbox("Log this application", value=True)
    with c2:
        more_bullets    = st.checkbox("Generate more bullet options")
    with c3:
        refresh_bullets = st.checkbox("Refresh only resume bullets")

    feedback = st.text_input("Optional: feedback or tone (e.g. friendly, persuasive)")
    st.caption("Tip: Professional, Friendly, Persuasive, Confident, Concise")

    # ⚡️ Single Generate button
    generate = st.button("✨ Generate AI Career Materials")

with col_help:
    st.markdown("**Quick Steps**")
    steps = [
        "Upload your current resume — drag & drop or browse.",
        "Paste the job description you want to apply to.",
        "Select options & hit **Generate AI Career Materials**.",
        "Not satisfied? Use feedback, expand bullets, or refresh.",
        "Download your personalized bullets, cover letter & message!"
    ]
    for i, step in enumerate(steps, 1):
        st.write(f"{i}. {step}")

# --- Bulk Job Descriptions ---
st.markdown("---")
st.markdown("## Bulk Job Descriptions")
st.caption("Download a CSV template and optionally upload multiple job descriptions.")
sample_csv = "Job Description\nJD #1\nJD #2\n"
st.download_button("Download Template CSV", data=sample_csv, file_name="job_description_template.csv")
bulk_file = st.file_uploader("Upload Bulk Job Descriptions CSV", type="csv")

# ✂️ Helpers to extract text
def extract_pdf(f):
    doc = fitz.open(stream=f.read(), filetype="pdf")
    return "".join(p.get_text() for p in doc)

def extract_docx(f):
    data = f.read()
    # reset for any future read
    f.seek(0)
    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs)

# --- Generation logic ---
if generate:
    # Validate
    if not resume_file:
        st.error("Please upload a resume.")
        st.stop()
    text = extract_pdf(resume_file) if resume_file.name.endswith(".pdf") else extract_docx(resume_file)

    # Collect JDs
    if bulk_file:
        dfj = pd.read_csv(bulk_file)
        jobs = dfj["Job Description"].dropna().tolist()
    elif job_desc:
        jobs = [job_desc]
    else:
        st.error("Please paste a job description or upload a CSV.")
        st.stop()

    jd0 = jobs[0]
    count = "five" if (more_bullets or refresh_bullets) else "two"
    fb    = f"Feedback: {feedback}\n" if feedback else ""

    prompt = f"""
{fb}You are an expert career coach AI. Using the resume below and the job description provided, return:
1. {count.capitalize()} tailored resume bullet points.
2. A personalized cover letter (needs to be 3 short paragraphs max).
3. A short outreach message to the hiring manager.

Resume:
{text}

Job Description:
{jd0}
"""
    resp = client.chat.completions.create(
        model=model,
        messages=[{"role":"user","content":prompt}],
        temperature=0.7
    )
   
    out = resp.choices[0].message.content.strip()

    # ── DEBUG CONSOLE ──
    st.code(out, language="markdown")

    # Improved section extraction based on actual headers
    bullets_section = cover_section = outreach_section = ""

    # Define the section headers as they appear in the response
    bullets_header = r"Tailored Resume Bullet Points:"
    cover_header = r"Personalized Cover Letter:"
    outreach_header = r"Short Outreach Message to Hiring Manager:"

    # Use regex to capture content between headers
    pattern = (
        rf"(?s){bullets_header}\s*(.*?)(?=\n{cover_header}|\n{outreach_header}|$)"
        rf"(?:\n{cover_header}\s*(.*?)(?=\n{outreach_header}|$))?"
        rf"(?:\n{outreach_header}\s*(.*))?"
    )
    match = re.search(pattern, out)
    if match:
        bullets_section = match.group(1).strip() if match.group(1) else ""
        cover_section = match.group(2).strip() if match.group(2) else ""
        outreach_section = match.group(3).strip() if match.group(3) else ""
    else:
        # Fallback: Split by headers if regex fails
        sections = re.split(rf"\n({bullets_header}|{cover_header}|{outreach_header})", out)
        for i in range(len(sections)):
            section = sections[i].strip()
            if i > 0 and sections[i-1] == bullets_header:
                bullets_section = section
            elif i > 0 and sections[i-1] == cover_header:
                cover_section = section
            elif i > 0 and sections[i-1] == outreach_header:
                outreach_section = section

    # Clean up the sections: Remove numbering (e.g., "1.", "2.") from the content
    def clean_section(text):
        # Remove leading numbers like "1.", "2.", etc., at the start of lines
        return re.sub(r"^\d+\.\s*", "", text, flags=re.MULTILINE).strip()

    bullets_text = clean_section(bullets_section)
    cover_text = clean_section(cover_section)
    outreach_text = clean_section(outreach_section)

    # Ensure sections are not empty
    if not bullets_text:
        bullets_text = "No resume bullet points generated. Please try again with a different job description or feedback."
    if not cover_text:
        cover_text = "No cover letter generated. Please try again with a different job description or feedback."
    if not outreach_text:
        outreach_text = "No outreach message generated. Please try again with a different job description or feedback."

    # Store in session
    st.session_state["bullets"] = bullets_text
    if not refresh_bullets:
        st.session_state["cover"] = cover_text
        st.session_state["outreach"] = outreach_text

    st.success("✅ Generated Successfully!")

# --- Display & Download Sections ---
if "bullets" in st.session_state:
    st.markdown("---")
    st.subheader("📌 Resume Bullets")
    st.markdown("**Tailored Resume Bullet Points:**")
    st.markdown(st.session_state["bullets"])
    buf1 = io.BytesIO()
    d1 = Document()
    d1.add_heading("Resume Bullets", 0)
    d1.add_paragraph(st.session_state["bullets"])
    d1.save(buf1)
    st.download_button("Download Resume Bullets", buf1.getvalue(), file_name="ResumeBullets.docx")

if "cover" in st.session_state:
    st.markdown("---")
    st.subheader("📜 Cover Letter")
    st.markdown(st.session_state["cover"])
    buf2 = io.BytesIO()
    d2 = Document()
    d2.add_heading("Cover Letter", 0)
    d2.add_paragraph(st.session_state["cover"])
    d2.save(buf2)
    st.download_button("Download Cover Letter", buf2.getvalue(), file_name="CoverLetter.docx")

if "outreach" in st.session_state:
    st.markdown("---")
    st.subheader("💬 Outreach Message")
    st.markdown("**Personalized Outreach Message:**")
    st.markdown(st.session_state["outreach"])
    buf3 = io.BytesIO()
    d3 = Document()
    d3.add_heading("Outreach Message", 0)
    d3.add_paragraph(st.session_state["outreach"])
    d3.save(buf3)
    st.download_button("Download Outreach Message", buf3.getvalue(), file_name="OutreachMessage.docx")

# --- Application Tracker & Log (no admin) ---
st.markdown("---")
st.subheader("📊 Application Tracker & Log")
if st.checkbox("Show Application Log"):
    try:
        df = pd.read_csv("application_log.csv", names=["Time","User","Job","Model","Preview"])
        df["Time"] = pd.to_datetime(df["Time"])
        df = df[df["User"].str.lower()==user_id]
        df = df.sort_values("Time", ascending=False)
        st.dataframe(df, use_container_width=True)
        csv_out = df.to_csv(index=False).encode("utf-8")
        st.download_button("Download Log as CSV", csv_out, file_name="application_log.csv")
    except FileNotFoundError:
        st.warning("No application log found yet. Generate your first application to start logging!")